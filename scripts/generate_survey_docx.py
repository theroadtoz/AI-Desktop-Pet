#!/usr/bin/env python3
"""将问卷 Markdown 转换为格式化的 .docx 文件。
格式标准：中文论文排版
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5


def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    if level == 0:
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    elif level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    return p


def body(text, indent=True, bold=False, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    return p


def meta_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(1)


def question(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    return p


def option(text, level=0):
    p = doc.add_paragraph()
    prefix = "    " * level
    run = p.add_run(prefix + text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    return p


def hint(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = None
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def divider():
    p = doc.add_paragraph()
    run = p.add_run('—' * 40)
    run.font.size = Pt(10)
    run.font.color.rgb = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════

heading('AI 桌面伙伴用户偏好调研问卷', level=0)
meta_line('版本 4.0 | 29 题 | 预计 6-8 分钟')
meta_line('目标人群：12-35 岁 Windows 电脑使用者')
meta_line('用途：收集潜在用户偏好，指导产品功能优先级与角色行为设计')
meta_line('')

# ══════════════════════════════════════════════
# 卷首语
# ══════════════════════════════════════════════

heading('卷首语', level=2)
body('你好！本问卷用于了解你对"桌面 AI 动漫角色"的使用偏好。所有答案仅用于毕业设计研究，无商业目的、不收集个人身份信息。')
body('你的第一反应就是最好的答案——不需要任何专业知识。')

divider()

# ══════════════════════════════════════════════
# 一、基本信息与使用习惯
# ══════════════════════════════════════════════

heading('一、基本信息与使用习惯', level=1)

question('Q1. 你的年龄段是？ [单选]')
option('A. 17 岁及以下')
option('B. 18-22 岁')
option('C. 23-26 岁')
option('D. 27-29 岁')
option('E. 30-39 岁')

question('Q2. 你目前的主要状态是？ [单选]')
option('A. 学生')
option('B. 全职工作')
option('C. 自由职业')
option('D. 其他')

question('Q3. 你对动漫 / 二次元 / ACG 内容的喜好程度？ [单选]')
option('A. 非常喜欢')
option('B. 比较喜欢')
option('C. 一般')
option('D. 不太感兴趣')
option('E. 完全不感兴趣')

question('Q4. 你平均每天使用电脑（台式机或笔记本）多长时间？ [单选]')
option('A. 不到 2 小时')
option('B. 2-4 小时')
option('C. 4-8 小时')
option('D. 8 小时以上')

question('Q5. 用电脑时，你有多大比例时间在运行全屏程序（如全屏游戏、全屏视频、全屏会议）？ [单选]')
option('A. 大部分时间全屏')
option('B. 一半一半')
option('C. 主要是窗口化多任务')
option('D. 不确定')

question('Q6. 你在电脑前能连续专注 1 小时以上不被任何事打断的频率？ [单选]')
option('A. 几乎每天都能')
option('B. 经常（一周 3-5 天）')
option('C. 偶尔（一周 1-2 天）')
option('D. 很少')

question('Q7. 你是否使用过 AI 对话 / 大语言模型产品（如 ChatGPT、豆包、文心一言、DeepSeek 等）？ [单选]')
option('A. 经常使用')
option('B. 偶尔使用')
option('C. 试过一两次')
option('D. 听说过但没用过')
option('E. 完全不了解')

question('Q8. 提到"桌面上的动漫角色"，你的第一反应是什么？请用几个词或一句话描述。 [开放题]')
hint('_____________________________________________________________')

divider()

# ══════════════════════════════════════════════
# 产品概念
# ══════════════════════════════════════════════

heading('产品概念', level=1)

body('桌面上的 Live2D 动漫角色，可以聊天。', bold=True, indent=False)
body('在你电脑桌面上显示一个会动的二次元角色。它不会挡住你的正常工作，可以拖拽到任何位置，点击它就能和它打字聊天。角色会根据对话内容改变表情和动作，像一个桌面上的小精灵。', indent=False)
body('请基于以上描述回答后续问题。', indent=False)

divider()

# ══════════════════════════════════════════════
# 二、认知与参照
# ══════════════════════════════════════════════

heading('二、认知与参照', level=1)

question('Q9. 在了解本产品之前，你使用过以下哪些类型的产品？ [多选]')
option('A. 桌面宠物 / 桌宠（如 Shimeji、Bongo Cat、Desktop Mate）')
option('B. AI 聊天 / 角色对话产品（如 Character.AI、豆包、星野）')
option('C. Live2D / 动态壁纸类产品')
option('D. 以上都没用过')

question('Q10. 上面描述的产品，和你之前用过的哪些东西最接近？ [多选]')
option('A. 以前养过的 QQ 宠物')
option('B. 手机上的虚拟助手（Siri / 小爱同学）')
option('C. 游戏里的 NPC 角色')
option('D. 电脑桌面小工具 / Widget')
option('E. Tamagotchi / 电子宠物')
option('F. 都差别很大，想不出接近的东西')
option('G. 其他（请填写）________')

divider()

# ══════════════════════════════════════════════
# 三、意愿与顾虑
# ══════════════════════════════════════════════

heading('三、意愿与顾虑', level=1)

question('Q11. 如果提供一个免费测试版，你有多大意愿安装并连续使用一周？ [单选]')
hint('[5 级量表]')
option('1 — 完全不愿意')
option('2 — 不太愿意')
option('3 — 一般 / 不确定')
option('4 — 比较愿意')
option('5 — 非常愿意')

question('Q12. 你对这类产品最大的担忧是什么？ [多选，限选 3 项]')
option('A. 影响电脑性能或变卡')
option('B. 太吵或太烦人，打扰工作')
option('C. AI 回复不稳定或重复')
option('D. 角色表情和动作不自然')
option('E. 担心聊天数据被上传或滥用')
option('F. 安装和配置太麻烦')
option('G. 担心收费或隐形消费')
option('H. 角色不好看或不喜欢')
option('I. 怕用几天就腻了')
option('J. 没有明显担忧')
option('K. 其他（请填写）________')

divider()

# ══════════════════════════════════════════════
# 四、角色行为设计
# ══════════════════════════════════════════════

heading('四、角色行为设计', level=1)

question('Q13. 你希望角色在你不主动找它的时候做什么？ [单选]')
option('A. 安静待着，几乎不动')
option('B. 偶尔做些小动作（眨眼、转头、伸懒腰）')
option('C. 根据时间段变化行为（早上精神、晚上犯困）')
option('D. 有存在感地活动（走动、做自己的事）')
option('E. 不确定')

question('Q14. 以下哪些场景你最不希望被角色打扰？ [多选]')
option('A. 全屏游戏时')
option('B. 视频会议 / 上网课时')
option('C. 全屏看视频时')
option('D. 专心写代码或写文档时')
option('E. 无所谓，不介意')
option('F. 其他（请填写）________')

question('Q15. 如果桌面角色有一个"存在感刻度"——从几乎注意不到到时刻提醒你它在，你偏好哪个位置？ [单选]')
hint('[5 级量表]')
option('1 — 几乎注意不到，偶尔瞥到就知道在')
option('2 — 比较安静，偶尔动一动')
option('3 — 中等，有明显存在感但不抢注意力')
option('4 — 比较活跃，经常能看到它在做各种事')
option('5 — 时刻提醒你它的存在，频繁互动')

question('Q16. 你希望角色主动和你说话吗？ [单选]')
option('A. 不希望，只有我点击它时才聊天')
option('B. 偶尔主动（如每天问候一次）')
option('C. 比较主动（在我空闲时主动发起话题）')
option('D. 不确定')

question('Q17. 你愿意让角色记住你的习惯和偏好吗？（如几点上班、喜欢聊什么话题、讨厌被什么方式打断） [单选]')
option('A. 愿意，会更贴心')
option('B. 可以，但要明确告诉我它记住了什么')
option('C. 不太愿意，感觉被监视')
option('D. 不确定')

divider()

# ══════════════════════════════════════════════
# 五、角色偏好
# ══════════════════════════════════════════════

heading('五、角色偏好', level=1)

question('Q18. 你希望角色主动和你互动的频率是？ [单选]')
option('A. 从不主动——只有我找它')
option('B. 每天 1 次（如早上问候）')
option('C. 每天 2-3 次')
option('D. 不固定频率，看情况')
option('E. 让我自己设置就好')
option('F. 不确定')

question('Q19. 你更能接受哪种视觉风格？ [单选]')
option('A. 日式动漫风格')
option('B. Q 版卡通 / 萌系')
option('C. 偏写实 2D')
option('D. 像素风')
option('E. 3D 角色')
option('F. 没有固定偏好')
option('G. 其他（请填写）________')

question('Q20. 你希望角色具备哪些人格特质？ [多选，限选 3 项]')
option('A. 元气活泼')
option('B. 温柔体贴')
option('C. 冷静理性')
option('D. 幽默风趣')
option('E. 调皮捣蛋')
option('F. 可靠稳重')
option('G. 呆萌可爱')
option('H. 有个性但情绪稳定、不崩人设')
option('I. 有分寸感，知道什么时候该安静')

question('Q21. 你希望角色默认待在桌面的哪个位置？ [单选]')
option('A. 屏幕右下角')
option('B. 屏幕左下角')
option('C. 跟随鼠标移动')
option('D. 坐在当前活动窗口的边缘')
option('E. 无所谓')

question('Q22. 你认为角色的大小应该大约是？ [单选]')
option('A. 拇指大小（约 2cm，几乎注意不到）')
option('B. 巴掌大小（约 5-8cm，不占空间）')
option('C. 手掌大小（约 10-12cm，有存在感但不挡视线）')
option('D. 更大（15cm 以上，想看清楚角色表情）')
option('E. 不确定')

divider()

# ══════════════════════════════════════════════
# 六、功能优先级
# ══════════════════════════════════════════════

heading('六、功能优先级', level=1)

question('Q23. 以下功能中请选择你认为最重要的 5 项。 [多选，限选 5 项]')
option('A. 透明悬浮，不阻挡操作桌面图标和窗口')
option('B. 可拖拽移动位置、可暂停或隐藏')
option('C. 根据对话内容改变表情')
option('D. 角色视线跟随鼠标')
option('E. 丰富自然的身体动作（不只表情，身体也会动）')
option('F. 点击 / 触摸角色时有反馈（摸头害羞、戳了生气）')
option('G. 低资源占用，不影响电脑性能')
option('H. 可以语音对话（角色能说话）')
option('I. 能记住之前的聊天内容')
option('J. 支持更换角色模型')
option('K. 坐在窗口边缘（吸附）')
option('L. 完全离线也能使用')

question('Q24. 以上功能中，你觉得哪一项最重要？ [单选]')
hint('（选项同 Q23）')

question('Q25. 以上功能中，你觉得哪一项最可以推迟到以后再实现？ [单选]')
hint('（选项同 Q23，增加"都不该推迟"）')

divider()

# ══════════════════════════════════════════════
# 七、隐私
# ══════════════════════════════════════════════

heading('七、隐私', level=1)

question('Q26. 如果这个桌面角色需要联网才能聊天（对话内容会发送给 AI 服务商处理），你的态度是？ [单选]')
option('A. 完全不能接受，必须支持离线使用')
option('B. 可以接受，但希望应用明确告知数据去了哪里')
option('C. 可以接受，不太在意')
option('D. 不确定')

divider()

# ══════════════════════════════════════════════
# 八、人口信息
# ══════════════════════════════════════════════

heading('八、人口信息', level=1)

question('Q27. 你的性别？ [单选]')
option('A. 男')
option('B. 女')
option('C. 不愿透露')

question('Q28. 你现在主要用的是哪类电脑？ [单选]')
option('A. Windows 台式机')
option('B. Windows 笔记本')
option('C. Mac 台式机 / MacBook')
option('D. 其他')

question('Q29. 如果你对这类产品还有其他想法、期待或担忧，请写在这里。 [开放题，选填]')
hint('_____________________________________________________________')
hint('_____________________________________________________________')

doc.add_paragraph()
body('问卷结束，感谢你的参与。', indent=False, bold=True, center=True)

# ─── 保存 ─────────────────────────────────────────────────

output_path = 'e:/Work-26/AI_Desktop_Pet/docs/reports/用户调研问卷-v4.0.docx'
doc.save(output_path)
print(f'[OK] Survey saved: {output_path}')
print(f'     Paragraphs: {len(doc.paragraphs)}')
