#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""从单一、已核验的 Markdown 源稿生成竞品分析正式报告。"""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "competitive-analysis-report.md"
OUTPUT = ROOT / "docs" / "reports" / "竞品分析报告.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "202124"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size=11, bold=False, color=BLACK, italic=False, mono=False):
    font = "Consolas" if mono else "Arial"
    east_asia = "宋体"
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Source Note" not in doc.styles:
        source_style = doc.styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source_style = doc.styles["Source Note"]
    source_style.font.name = "Arial"
    source_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    source_style.font.size = Pt(9)
    source_style.font.color.rgb = RGBColor.from_string(MID_GRAY)
    source_style.paragraph_format.space_before = Pt(4)
    source_style.paragraph_format.space_after = Pt(4)
    source_style.paragraph_format.line_spacing = 1.0


def add_inline_text(paragraph, text, size=11, color=BLACK):
    token_re = r"(\*\*.*?\*\*|`.*?`|https?://\S+)"
    for part in re.split(token_re, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run_font(paragraph.add_run(part[2:-2]), size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            set_run_font(paragraph.add_run(part[1:-1]), size=size, color=color, mono=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_title_page(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for _ in range(5):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("市场研究 / 竞品核验"), size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("AI 桌面伙伴竞品分析报告"), size=27, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(
        subtitle.add_run("从“市场空白”叙事转向可验证的长期驻留体验"),
        size=14,
        color=BLUE,
    )

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    set_table_geometry(meta, [2200, 7160])
    values = [
        ("核验日期", "2026-06-11"),
        ("研究对象", "Windows 桌面伙伴、AI 角色伴侣、桌宠生态与开源参考"),
        ("证据口径", "官网、Steam、官方 GitHub 与官方文档；评论仅作个案信号"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        set_run_font(p.add_run(label), size=9.5, bold=True, color=NAVY)
        p = row.cells[1].paragraphs[0]
        set_run_font(p.add_run(value), size=9.5)

    doc.add_paragraph()
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(22)
    set_run_font(
        notice.add_run("本报告取代旧版“第一、唯一、完全空白”等未经证实结论。"),
        size=10,
        color=MID_GRAY,
        italic=True,
    )
    doc.add_page_break()


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("AI 桌面伙伴竞品分析 | "), size=8.5, color=MID_GRAY)
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_run_font(run, size=8.5, color=MID_GRAY)


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, level_text, suffix, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_numbered_list(doc, items):
    num_id = create_decimal_numbering(doc)
    for text in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.1
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        add_inline_text(p, text)


def parse_table(lines):
    rows = []
    for line in lines:
        values = [item.strip() for item in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        rows.append(values)
    return rows


def choose_widths(rows):
    columns = max(len(row) for row in rows)
    if columns == 3:
        return [1500, 3830, 4030]
    if columns == 4:
        return [1350, 2550, 2600, 2860]
    if columns == 5:
        return [1150, 1750, 2450, 1500, 2510]
    return [TABLE_WIDTH_DXA // columns] * columns


def add_table(doc, rows):
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = choose_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for row_index, (row, values) in enumerate(zip(table.rows, rows)):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value = values[col_index] if col_index < len(values) else ""
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if row_index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_text(p, value, size=8.5, color=WHITE)
                for run in p.runs:
                    run.font.bold = True
            else:
                add_inline_text(p, value, size=8.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_markdown(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    skipped_title = False

    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            if text and not text.startswith(("版本：", "研究范围：", "证据原则：")):
                p = doc.add_paragraph(style="Source Note")
                add_inline_text(p, text, size=9, color=MID_GRAY)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading = re.match(r"^(#{2,4})\s+(.*)", line)
        if heading:
            level = len(heading.group(1)) - 1
            doc.add_paragraph(heading.group(2), style=f"Heading {level}")
        elif re.match(r"^\d+\.\s+", line):
            items = []
            while index < len(lines):
                match = re.match(r"^\d+\.\s+(.*)", lines[index].rstrip())
                if not match:
                    break
                items.append(match.group(1))
                index += 1
            add_numbered_list(doc, items)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(6)
            add_inline_text(p, line[2:])
        elif re.match(r"^https?://", line) or "：https://" in line:
            p = doc.add_paragraph(style="Source Note")
            add_inline_text(p, line, size=9, color=MID_GRAY)
        else:
            p = doc.add_paragraph()
            add_inline_text(p, line)
        index += 1


def build_document():
    doc = Document()
    configure_styles(doc)
    add_title_page(doc)
    add_markdown(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        add_footer(section)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"竞品报告已生成: {OUTPUT}")


if __name__ == "__main__":
    build_document()
