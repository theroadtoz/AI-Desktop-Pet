#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成竞品联网核验过程记录。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "process" / "竞品分析阶段_过程记录.docx"
BLUE = "1F4E78"
LIGHT = "EAF2F8"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def style(run, size=10.5, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def para(doc, text="", size=10.5, bold=False, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style(p.add_run(text), size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    sizes = {0: 19, 1: 14, 2: 12}
    p = para(doc, text, size=sizes[level], bold=True, center=level == 0)
    p.paragraph_format.space_before = Pt(10 if level else 0)
    p.paragraph_format.space_after = Pt(9)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, value in enumerate(headers):
        shade(t.rows[0].cells[index], BLUE)
        p = t.rows[0].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style(p.add_run(value), size=9, bold=True, color=RGBColor(255, 255, 255))
    for row_index, values in enumerate(rows):
        cells = t.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                shade(cells[index], LIGHT)
            style(cells[index].paragraphs[0].add_run(str(value)), size=8.8)
    para(doc, "")


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    heading(doc, "AI 桌面伙伴项目", 0)
    heading(doc, "竞品联网核验过程记录", 0)
    para(doc, "版本 2.0 | 核验日期 2026-06-09", center=True)

    heading(doc, "一、核验目标", 1)
    para(doc, "检查2026-05-31竞品研究中的产品事实、市场判断、商业数字、技术路线和研究可复现性，并区分官方事实、第三方估算、用户评论和研究者推断。")

    heading(doc, "二、并行审查分工", 1)
    table(
        doc,
        ["审查方向", "主要任务"],
        [
            ["竞品事实", "核对名称、开发商、发布日期、价格、平台、功能和动态数据"],
            ["市场与商业", "核对市场空白、买断验证、订阅疲劳、MAU和商业数字"],
            ["技术路线", "核对Electron、Live2D、Pixi、点击穿透和恢复API"],
            ["问卷方法", "核对版本、诱导性、量表、抽样、统计和隐私伦理"],
            ["独立红队", "寻找反例、选择性引用、确认偏误和错误立项风险"],
        ],
    )

    heading(doc, "三、来源优先级", 1)
    table(
        doc,
        ["优先级", "来源", "使用规则"],
        [
            ["1", "Steam官方页、GitHub官方仓库、产品官网、官方文档", "产品功能、价格、版本和技术事实的主要依据"],
            ["2", "监管机构、公司年报、官方隐私政策", "法律事件、交易性质和数据使用依据"],
            ["3", "可靠媒体和行业数据", "必须标注发布日期、口径和属于第三方估算"],
            ["4", "Steam评论、Reddit、B站等用户内容", "只能做定性样本，必须披露抽样方法"],
            ["禁止", "无链接的AI生成数字或动机判断", "不得写成事实"],
        ],
    )

    heading(doc, "四、主要纠正", 1)
    table(
        doc,
        ["旧结论", "核验结果", "处理"],
        [
            ["市场完全空白、第一、唯一", "已存在AI Desktop Pet、Deskie、Moe.AI等直接竞品", "删除"],
            ["Desktop Mate由MegaRads开发", "Steam显示开发商为infiniteloop", "纠正"],
            ["Project N.E.K.O.是买断制且已支持移动端", "当前为免费抢先体验，移动/VR仅为方向", "纠正"],
            ["AIRI采用Tauri", "当前官方桌面端采用Electron", "纠正"],
            ["约18元上架证明买断可行", "缺乏销量、留存、退款和盈利数据", "降级为待验证假设"],
            ["Google向Character.AI投资27亿美元覆盖亏损", "交易主要涉及技术许可和人才，资金用途推断无依据", "纠正"],
            ["数据永不出电脑", "云端API会接收当前请求", "改为分模式说明"],
            ["Electron 30为安全锁版", "Electron 30已停止支持", "改为验证当前受支持版本"],
        ],
    )

    heading(doc, "五、研究可复现性要求", 1)
    para(doc, "今后的竞品记录为每条事实保存URL、访问日期、页面存档、事实/推断标签和数据口径。")
    para(doc, "评分矩阵不得将未开发产品的目标分数与竞品实测分数并列。优先改用步骤数、启动时间、资源占用、平台支持和可重复测试结果。")
    para(doc, "用户评论分析必须披露平台、日期范围、语言、样本量、纳入规则和编码方法。")

    heading(doc, "六、关键官方来源", 1)
    sources = [
        "AI Desktop Pet: https://store.steampowered.com/app/4227700/AI_Desktop_Pet/",
        "Deskie: https://store.steampowered.com/app/3540710/Deskie/",
        "Moe.AI: https://store.steampowered.com/app/4159250/MoeAI/",
        "Project N.E.K.O.: https://store.steampowered.com/app/4099310/Project_NEKO/",
        "Desktop Mate: https://store.steampowered.com/app/3301060/Desktop_Mate/",
        "Open-LLM-VTuber: https://github.com/Open-LLM-VTuber/Open-LLM-VTuber",
        "AIRI: https://github.com/moeru-ai/airi",
        "Electron窗口交互: https://www.electronjs.org/docs/latest/tutorial/custom-window-interactions",
        "Live2D Cubism SDK: https://www.live2d.com/en/sdk/download/web/",
        "Character.AI模型训练说明: https://policies.character.ai/model-training",
        "意大利数据保护机构: https://www.gpdp.it/garante/doc.jsp?ID=10132048",
    ]
    for source in sources:
        p = doc.add_paragraph(style="List Bullet")
        style(p.add_run(source), size=9.5)

    heading(doc, "七、修订结果", 1)
    para(doc, "已重写competitive-comparison.md和competitive-strategy.md，修订CONTEXT.md、MASTER_PLAN.md和task-list.md，并重新生成正式竞品报告。")
    para(doc, "旧版review-report.md保留为历史文档，但已添加失效提示，不能继续作为当前实施依据。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"竞品过程记录已生成: {OUTPUT}")


if __name__ == "__main__":
    build_document()
