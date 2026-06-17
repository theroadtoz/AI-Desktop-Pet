#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成中性、可执行的 AI 桌面伙伴探索问卷。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "reports" / "用户调研分析问卷.docx"

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
GRAY = RGBColor(0x66, 0x66, 0x66)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def set_font(run, size=11, bold=False, color=None, cn="宋体", en="Arial"):
    run.font.name = en
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def paragraph(doc, text="", size=11, bold=False, center=False, after=5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {0: 20, 1: 15, 2: 12}
    p = paragraph(
        doc,
        text,
        size=sizes[level],
        bold=True,
        center=level == 0,
        after=10 if level else 16,
    )
    if level:
        p.paragraph_format.keep_with_next = True
    return p


def question(doc, number, text, kind):
    p = paragraph(doc, f"Q{number}. {text}  [{kind}]", size=11, bold=True, after=4)
    p.paragraph_format.keep_with_next = True


def options(doc, values):
    for index, value in enumerate(values):
        p = paragraph(doc, f"    {chr(65 + index)}. {value}", size=10, after=2)
        p.paragraph_format.left_indent = Cm(0.35)
    paragraph(doc, "", size=4, after=2)


def response_lines(doc, count=2):
    for _ in range(count):
        paragraph(doc, "__________________________________________________________________", size=9, color=GRAY)


def scale(doc, left, right):
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, cell in enumerate(table.rows[0].cells, start=1):
        cell.width = Cm(2.1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(str(index)), size=10, bold=True)
        if index % 2 == 0:
            shade(cell, LIGHT_BLUE)
    labels = [left, "", "一般/不确定", "", right]
    for index, (cell, label) in enumerate(zip(table.rows[1].cells, labels)):
        cell.width = Cm(2.1)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=8, color=GRAY)
    paragraph(doc, "", size=3)


def callout(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "EDF4F8")
    p = cell.paragraphs[0]
    set_font(p.add_run(title), size=11, bold=True, color=RGBColor(0x1F, 0x4E, 0x78))
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(line), size=10)
    paragraph(doc, "", size=4)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    heading(doc, "AI 桌面伙伴概念探索问卷", 0)
    paragraph(doc, "版本 3.0 | 31 题 | 预计 8-10 分钟", center=True, color=GRAY)

    callout(
        doc,
        "研究说明",
        [
            "本研究用于了解 Windows 用户对 AI 桌面角色的真实使用需求，不存在正确答案。",
            "问卷不收集姓名和联系方式。试玩报名应使用独立表单，不能与问卷回答关联。",
            "你可以跳过不愿回答的问题，也可以随时退出。",
            "本问卷只用于探索性研究，结果不代表全部 Windows 用户。",
        ],
    )

    heading(doc, "一、资格与当前行为", 1)
    question(doc, 1, "你的年龄是？", "单选")
    options(doc, ["17岁及以下（结束问卷）", "18-22岁", "23-26岁", "27-29岁", "30-39岁", "40岁及以上"])

    question(doc, 2, "你目前是否经常使用 Windows 电脑？", "单选")
    options(doc, ["是，每周5天及以上", "是，每周2-4天", "偶尔使用", "基本不使用（结束问卷）"])

    question(doc, 3, "在通常的一天里，你使用电脑多长时间？", "单选")
    options(doc, ["少于2小时", "2至不足4小时", "4至不足8小时", "8小时及以上"])

    question(doc, 4, "你对桌面宠物或桌面角色的使用情况是？", "单选")
    options(doc, ["正在使用", "以前使用过，现在不用了", "听说过但没用过", "从未听说过"])

    question(doc, 5, "你对 AI 角色聊天产品的使用情况是？", "单选")
    options(doc, ["正在使用", "以前使用过，现在不用了", "听说过但没用过", "从未听说过"])

    question(doc, 6, "过去30天，你使用桌面宠物或AI角色聊天的频率是？", "单选")
    options(doc, ["每天", "每周数次", "每周一次左右", "少于每周一次", "没有使用"])

    question(doc, 7, "如果你停止使用过类似产品，最主要原因是什么？", "单选，可跳过")
    options(doc, ["新鲜感消失", "影响电脑性能", "互动重复或角色不自然", "配置麻烦", "价格或持续费用", "隐私担忧", "会打扰工作", "其他", "没有停止使用过"])

    heading(doc, "二、概念前的自发需求", 1)
    question(doc, 8, "在没有看到产品介绍前，你理想中的桌面伙伴应解决什么问题或带来什么感受？", "开放题")
    response_lines(doc, 3)

    question(doc, 9, "什么情况最可能让你卸载或关闭一款桌面伙伴？", "开放题")
    response_lines(doc, 2)

    heading(doc, "三、中性概念测试", 1)
    callout(
        doc,
        "产品概念",
        [
            "这是一款 Windows 桌面应用。一个 Live2D 角色可以置顶显示、拖动和进行基础鼠标互动。",
            "点击角色后可以打开文字聊天。系统根据回复选择表情和动作。",
            "首版计划只提供一个官方角色，语音、长期记忆和模型导入可能不会同时提供。",
            "云端模型通常回复质量较高，但需要联网，并会把当前请求发送给模型供应商。",
            "本地模型可以不把对话发送到外部服务，但需要下载模型并占用更多存储、内存和算力。",
            "应用是否收费、价格和默认模型尚未确定。",
        ],
    )

    question(doc, 10, "上述概念是否容易理解？", "5点评分")
    scale(doc, "完全不理解", "完全理解")

    question(doc, 11, "整体上，这个概念对你有多大吸引力？", "5点评分")
    scale(doc, "完全没有吸引力", "非常有吸引力")

    question(doc, 12, "如果今天提供安全的测试版，你在未来30天安装试用的可能性有多大？", "5点评分")
    scale(doc, "肯定不会", "肯定会")

    question(doc, 13, "你对这个概念最大的顾虑是什么？", "最多选3项")
    options(doc, ["影响性能或续航", "打扰工作", "AI回复错误或人设不稳定", "动作或表情不自然", "隐私和数据使用", "安装或模型配置复杂", "价格或持续费用", "角色风格不喜欢", "容易失去新鲜感", "没有明显顾虑", "其他"])

    heading(doc, "四、MVP 功能取舍", 1)
    features = [
        "透明置顶且不阻挡正常鼠标操作",
        "拖动、隐藏、暂停和托盘控制",
        "文字聊天",
        "表情随对话变化",
        "身体动作随状态变化",
        "视线、点击和触摸反馈",
        "低资源占用",
        "语音朗读回复",
        "跨会话记忆",
        "更换或导入角色模型",
        "窗口边缘吸附或坐在窗口上",
        "完全本地模型模式",
    ]
    question(doc, 14, "如果首版最多实现5项，你会选择哪些？", "最多选5项")
    options(doc, features)

    question(doc, 15, "在上题中，你认为最重要的一项是什么？", "填写一项")
    response_lines(doc, 1)

    question(doc, 16, "在上题中，你认为最可以推迟的一项是什么？", "填写一项")
    response_lines(doc, 1)

    question(doc, 17, "你希望角色主动发起互动的频率是？", "单选")
    options(doc, ["从不主动", "每天最多1次", "每天2-3次", "按早中晚等固定场景", "由我设置具体规则", "不确定"])

    question(doc, 18, "你能接受哪些视觉风格？", "多选")
    options(doc, ["日式二次元", "Q版卡通", "较写实2D", "像素风", "3D角色", "没有固定偏好", "其他"])

    question(doc, 19, "你更重视角色的哪些人格特征？", "最多选3项")
    options(doc, ["主动程度可控", "情绪表达丰富", "回复简洁", "幽默", "温柔支持", "冷静理性", "调皮活泼", "人格长期稳定", "可自行调整"])

    heading(doc, "五、技术与隐私权衡", 1)
    question(doc, 20, "在以下模式中，你最倾向哪一种？", "单选")
    options(doc, ["本地模型：更私密，但下载大、占资源", "云端模型：质量和速度可能更好，但请求发送给供应商", "自带API Key：自己承担模型费用，但需要配置", "由产品提供有限云额度，超出后付费", "不确定"])

    question(doc, 21, "你愿意为了使用更好的云端模型，自行填写 API Key 吗？", "单选")
    options(doc, ["愿意", "可以接受，但需要详细引导", "不愿意", "不了解API Key"])

    question(doc, 22, "为了完全本地运行，你最多能接受多大的额外模型下载？", "单选")
    options(doc, ["不接受额外下载", "2GB以内", "2-5GB", "5-10GB", "10GB以上", "不确定"])

    question(doc, 23, "隐私模式会在多大程度上影响你的购买或使用决定？", "5点评分")
    scale(doc, "完全不影响", "决定性影响")

    heading(doc, "六、商业模式探索", 1)
    question(doc, 24, "你更能接受哪种付费方式？", "单选")
    options(doc, ["基础应用一次买断", "应用免费，角色或语音包单独买断", "免费基础版，高级云模型订阅", "完全免费，但自行配置模型/API", "只考虑完全免费产品", "不确定"])

    question(doc, 25, "一次买断价格低于多少元时，你会怀疑产品质量或后续维护？", "填写金额")
    response_lines(doc, 1)

    question(doc, 26, "一次买断价格低于多少元时，你会认为比较划算？", "填写金额")
    response_lines(doc, 1)

    question(doc, 27, "一次买断价格高于多少元时，你会开始觉得贵，但仍可能购买？", "填写金额")
    response_lines(doc, 1)

    question(doc, 28, "一次买断价格高于多少元时，你基本不会购买？", "填写金额")
    response_lines(doc, 1)

    question(doc, 29, "哪些条件会让你更愿意实际安装测试版？", "最多选3项")
    options(doc, ["有完整演示视频", "有明确性能数据", "有隐私数据流说明", "有喜欢的官方角色", "可以免费试玩", "支持退款", "朋友或创作者推荐", "其他"])

    heading(doc, "七、背景信息", 1)
    question(doc, 30, "你目前的主要状态是？", "单选")
    options(doc, ["学生", "全职工作", "兼职工作", "自由职业", "待业或求职", "其他"])

    question(doc, 31, "你对动漫、游戏角色和虚拟主播等 ACG 内容的总体态度是？", "单选")
    options(doc, ["非常喜欢", "比较喜欢", "一般", "不太喜欢", "非常不喜欢"])

    heading(doc, "问卷结束", 0)
    paragraph(doc, "感谢参与。试玩报名请通过独立表单收集，避免联系方式与问卷答案关联。", center=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"问卷已生成: {OUTPUT}")


if __name__ == "__main__":
    build_document()
