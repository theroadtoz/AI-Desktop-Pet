#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成与 v3.0 问卷一致的研究过程记录。"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "process" / "用户调研阶段_过程记录.docx"
BLUE = "1F4E78"
LIGHT = "D9EAF7"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    props.append(node)


def font(run, size=11, bold=False, color=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def para(doc, text="", size=11, bold=False, center=False, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(text), size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    size = {0: 20, 1: 15, 2: 12}[level]
    return para(doc, text, size=size, bold=True, center=level == 0, after=11)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for index, text in enumerate(headers):
        cell = t.rows[0].cells[index]
        shade(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        font(run, size=9, bold=True, color=RGBColor(255, 255, 255))
    for row_index, values in enumerate(rows):
        cells = t.add_row().cells
        for col_index, value in enumerate(values):
            if row_index % 2:
                shade(cells[col_index], LIGHT)
            p = cells[col_index].paragraphs[0]
            font(p.add_run(str(value)), size=9)
    para(doc, "", size=3)
    return t


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    heading(doc, "AI 桌面伙伴用户研究", 0)
    heading(doc, "问卷设计与执行记录", 0)
    para(doc, "版本 3.0 | 与31题正式问卷一致 | 2026-06-09", center=True)

    heading(doc, "一、研究定位", 1)
    para(doc, "本研究定位为探索性概念测试，不用于估计全部 Windows 用户的市场需求比例，也不直接证明产品商业可行。")
    para(doc, "研究目标是发现：目标用户当前行为、概念吸引力、主要顾虑、MVP功能取舍、本地与云端模式偏好，以及初步价格区间。")

    heading(doc, "二、旧版问题与修订", 1)
    table(
        doc,
        ["旧版问题", "风险", "v3.0修订"],
        [
            ["过程记录称33题，实际问卷26题且封面写22题", "研究不可复现", "统一为7部分、31题"],
            ["使用“小生命、她懂你”等宣传语言", "正向启动和诱导", "改为中性描述，并同时呈现限制"],
            ["宣称数据不上传，但默认规划云端API", "事实错误和合规风险", "明确本地与云端数据流"],
            ["没有商业模式和价格题", "无法支持买断定价", "加入模式选择和PSM四问"],
            ["30份便利样本计划做复杂交叉分析", "统计不稳定且不可外推", "分阶段研究并限制结论"],
            ["声称匿名但同表收集联系方式", "不是真正匿名", "联系方式改由独立表单收集"],
            ["全部功能做正向重要性评分", "天花板效应，缺少取舍", "首版最多选5项，再选最重要和最可推迟"],
        ],
    )

    heading(doc, "三、问卷结构", 1)
    table(
        doc,
        ["部分", "题号", "目的"],
        [
            ["资格与当前行为", "Q1-Q7", "筛选成年人和Windows用户，测量真实使用行为"],
            ["概念前自发需求", "Q8-Q9", "在产品介绍前收集未被提示的需求与卸载原因"],
            ["中性概念测试", "Q10-Q13", "测量理解度、吸引力、试用可能性和顾虑"],
            ["MVP功能取舍", "Q14-Q19", "强迫资源约束下的功能选择和交互偏好"],
            ["技术与隐私权衡", "Q20-Q23", "比较本地、云端、BYOK和下载成本"],
            ["商业模式探索", "Q24-Q29", "模式偏好、PSM价格问题和真实试玩条件"],
            ["背景信息", "Q30-Q31", "用于有限的探索性分组描述"],
        ],
    )

    heading(doc, "四、抽样与执行方案", 1)
    para(doc, "阶段1：5-8人认知访谈。逐题观察理解偏差、术语问题和选项遗漏。")
    para(doc, "阶段2：10-15人预测试。检查完成时间、跳题逻辑、题序和数据质量。")
    para(doc, "阶段3：100-200人的探索性问卷。建议按ACG兴趣高低、是否用过桌宠、学生/在职设置配额；每个关键分组尽量保留30-50人。")
    para(doc, "招募仍属于非概率样本。报告只能描述本次样本，不得使用“代表年轻人”“市场中有多少比例”等外推语言，也不得计算传统概率样本误差范围。")

    heading(doc, "五、数据质量规则", 1)
    table(
        doc,
        ["检查项", "处理原则"],
        [
            ["极短完成时间", "作为风险信号，不单独判无效"],
            ["逻辑矛盾", "结合资格题和后续行为题人工复核"],
            ["开放题乱码或无意义内容", "可作为无效依据"],
            ["重复设备或账号", "保留一份并记录规则"],
            ["长串同值", "只作为辅助指标，不能自动删除"],
            ["排他选项冲突", "在线问卷设置互斥逻辑"],
        ],
    )

    heading(doc, "六、分析计划", 1)
    para(doc, "描述性分析：报告各题人数和比例，并同时报告有效样本数。")
    para(doc, "分组分析：只比较预先指定的少量分组；样本过小时仅描述，不做显著性推断。")
    para(doc, "MVP排序：统计最多选5项的入选率、最重要项和最可推迟项。条件允许时，后续研究可改用MaxDiff。")
    para(doc, "PSM：Q25-Q28用于探索价格边界。结果只能作为价格实验输入，不能替代真实购买、愿望单、试玩转化和退款数据。")
    para(doc, "IPA不适用于当前概念问卷。只有原型可实际体验并同时测量重要性与满意度后，才考虑使用。")

    heading(doc, "七、隐私与伦理", 1)
    para(doc, "问卷限定18岁及以上；不在正式问卷中收集姓名、邮箱、微信或QQ。")
    para(doc, "试玩报名使用独立表单，不使用可关联的响应ID。")
    para(doc, "发布问卷时应说明研究目的、数据用途、保存期限、访问人员和退出方式。")

    heading(doc, "八、报告披露要求", 1)
    para(doc, "最终报告必须披露目标总体、招募渠道、调查日期、问卷全文、样本构成、完成率、排除规则、题序、是否随机化以及研究局限。")
    para(doc, "所有结论区分：本次样本观察、用户原话、研究者推断和待真实行为验证的假设。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"过程记录已生成: {OUTPUT}")


if __name__ == "__main__":
    build_document()
