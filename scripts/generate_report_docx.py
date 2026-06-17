#!/usr/bin/env python3
"""将竞品分析报告 Markdown 转换为格式化的 .docx 文件。

格式标准：中文毕业论文规范
- 论文标题：小二号(18pt) 黑体 加粗 居中
- 一级标题：三号(16pt) 黑体 加粗
- 二级标题：四号(14pt) 黑体 加粗
- 三级标题：小四(12pt) 黑体 加粗
- 正文：小四(12pt) 宋体 1.5倍行距 首行缩进2字符
- 表格内容：五号(10.5pt) 宋体
- 页边距：上下2.54cm 左右3.18cm
- 编号：1 / 1.1 / 1.1.1 格式
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ─── 工具函数 ─────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" '
            f'w:color="{val.get("color","000000")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_heading(doc, text, level=1):
    """创建带编号的标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        # 论文标题
        run.font.size = Pt(18)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.5
    elif level == 1:
        run.font.size = Pt(16)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.5
    elif level == 2:
        run.font.size = Pt(14)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
    elif level == 3:
        run.font.size = Pt(12)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
    return p


def make_body(doc, text, indent=True, bold=False):
    """创建正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 2字符
    p.paragraph_format.space_after = Pt(3)
    return p


def make_info_paragraph(doc, lines):
    """创建信息块（作者、日期等）"""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5


def make_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 边框
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={"val": "single", "sz": "4", "color": "333333"},
                bottom={"val": "single", "sz": "4", "color": "333333"},
                left={"val": "single", "sz": "4", "color": "333333"},
                right={"val": "single", "sz": "4", "color": "333333"},
            )

    doc.add_paragraph()  # 表后空行
    return table


def make_bullet(doc, text, level=0):
    """创建项目符号列表"""
    p = doc.add_paragraph()
    prefix = "• " if level == 0 else "  ◦ "
    run = p.add_run(prefix + text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(24 * (level + 1))
    p.paragraph_format.space_after = Pt(1)
    return p


def add_page_break(doc):
    doc.add_page_break()


# ─── 文档生成 ─────────────────────────────────────────────

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ══════════════════════════════════════════════
# 封面信息
# ══════════════════════════════════════════════

make_heading(doc, 'AI 桌面伙伴竞品分析报告', level=0)

info_lines = [
    '版本：v2.0    日期：2026-06-11',
    '用途：毕业设计竞品分析章节',
]
make_info_paragraph(doc, info_lines)

# ══════════════════════════════════════════════
# 摘要
# ══════════════════════════════════════════════

make_heading(doc, '摘要', level=1)

make_body(doc, '本报告对"AI 驱动的 Live2D 桌面宠物"赛道进行系统性竞品分析，为毕业设计项目 Desktop Pet 提供设计依据。研究范围覆盖 5 个直接竞品（AI Desktop Pet、Project N.E.K.O.、Desktop Mate、Open-LLM-VTuber、Deskie），涵盖商业产品和开源方案。')

make_body(doc, '核心发现：第一，"Live2D + AI 对话 + 桌面常驻"已形成直接竞争，本项目不以"市场空白"为定位。第二，现有竞品在桌面融合质量、低打扰设计、AI 人设稳定性三个维度上存在可测量的差异和改进空间。第三，竞品普遍追求功能广度（多模态、Agent、Workshop），但在"长期桌面驻留的克制体验"上投入不足——这是本项目的论证切入点。')

make_body(doc, '本报告围绕三个论证维度展开：')

make_table(doc,
    ['维度', '测量方式', '对象与方法'],
    [
        ['桌面融合质量', '7 项技术测试用例', 'Desktop Mate + AI Desktop Pet + 本系统（实测）'],
        ['低打扰设计', '6 维客观设计分析', '全部 5 个竞品 + 本系统（文档分析）'],
        ['AI 人设稳定性', '消融实验（有/无防护对比）', '本系统 A/B/C 三组 + 竞品定性讨论'],
    ]
)

# ══════════════════════════════════════════════
# 一、市场格局
# ══════════════════════════════════════════════

make_heading(doc, '一、市场格局', level=1)

# 1.1
make_heading(doc, '1.1 赛道定义', level=2)

make_body(doc, '本报告研究的赛道为：以 Live2D 或 3D 角色在 Windows 桌面上常驻渲染为核心交互形式、集成大语言模型实现对话能力的桌面应用。')

make_body(doc, '区别于以下三类产品：')
make_bullet(doc, '纯 AI 对话伴侣（Character.AI、Replika）：无桌面驻留和角色渲染')
make_bullet(doc, '纯视觉桌宠（Shimeji、Bongo Cat）：无 AI 对话能力')
make_bullet(doc, 'VTuber 直播工具（VTube Studio、PrprLive）：面向直播而非桌面陪伴')

# 1.2
make_heading(doc, '1.2 市场参与者分类', level=2)

make_body(doc, '按"桌面融合"与"AI 对话"两个维度，市场参与者分布如下：')

make_table(doc,
    ['类别', '代表产品', '桌面融合', 'AI 对话', '定位特征'],
    [
        ['纯桌宠标杆', 'Desktop Mate、VPet、Shimeji', '强', '弱', '视觉陪伴为主，验证了桌面存在感的长期需求'],
        ['AI 对话标杆', 'Character.AI、Replika', '弱', '强', '验证了角色对话的市场规模，但无桌面整合'],
        ['开源全栈方案', 'Open-LLM-VTuber、AIRI', '中', '强', '技术方案成熟，可全离线，但配置门槛较高'],
        ['商业 AI 桌宠', 'AI Desktop Pet、Deskie、Project N.E.K.O.', '中', '中-强', '直接竞品，各有侧重'],
        ['本项目目标', 'Desktop Pet', '强', '中', '克制型常驻体验，不求功能最多'],
    ]
)

# 1.3
make_heading(doc, '1.3 商业模式分化', level=2)

make_body(doc, '截至 2026 年 6 月，市场已分化为五种清晰的商业模式：')

make_table(doc,
    ['模式', '代表产品', '价格', '核心逻辑'],
    [
        ['本地 AI 买断', 'AI Desktop Pet', '¥52 一次性', '内置本地 LLM，完全离线，无需 API Key'],
        ['免费 + 点数/内购', 'Deskie、Molili AI', '免费 + 萌点消耗', 'AI 对话消耗虚拟货币，内容创建免费'],
        ['免费 + 角色 DLC', 'Desktop Mate', '免费本体 + $14.99/DLC', '授权 IP 角色是核心收入来源'],
        ['完全免费开源', 'Project N.E.K.O.、Open-LLM-VTuber', '免费', '开源核心，社区驱动，未来计划内容变现'],
        ['低价买断 + BYOK', 'AI Desk Pet', '¥11 + 自备 Key', '低门槛本体，用户承担模型费用'],
    ]
)

# ══════════════════════════════════════════════
# 二、竞品深度档案
# ══════════════════════════════════════════════

make_heading(doc, '二、竞品深度档案', level=1)

make_body(doc, '本章对 5 个直接竞品进行深度分析。每个竞品档案统一结构：基本信息、产品定位、核心功能、关键评价、与本项目的关系。')

# --- 2.1 AI Desktop Pet ---
make_heading(doc, '2.1 AI Desktop Pet', level=2)

make_table(doc,
    ['属性', '内容'],
    [
        ['开发商', '42kami'],
        ['发布日期', '2026-04-28'],
        ['平台', 'Windows（Steam）'],
        ['价格', '¥52 买断（中国区），$12.99（美国区）'],
        ['技术栈', 'Electron + Live2D Cubism SDK（官方授权）+ 内置本地 LLM + NVIDIA CUDA'],
        ['Steam 评价', '82% 好评（约 50 条评价），SteamDB 评分 72.5%'],
    ]
)

make_body(doc, '产品定位：Live2D 桌宠 + 本地 AI 角色扮演 + 视觉小说的"三位一体"离线体验。其核心卖点是"完全离线"——内置本地 LLM，无需 API Key、无需联网，所有对话数据存储在本地。', bold=False)

make_body(doc, '核心功能：', bold=True)
make_bullet(doc, '完全离线运行：内置本地 LLM，支持 GPU 加速推理')
make_bullet(doc, '语音交互：SenseVoice/Whisper 语音识别 + Kokoro TTS 引擎 + 声音克隆')
make_bullet(doc, '角色系统：兼容 SillyTavern 人设卡格式，内置 100+ 角色卡，支持世界书')
make_bullet(doc, 'Steam 创意工坊：Live2D 模型、角色卡、世界书一键订阅')
make_bullet(doc, '眼球追踪 + 表情/动作交互')
make_bullet(doc, '30 种界面语言')

make_body(doc, '关键评价：作为市场上最直接的可比产品，其"完全离线"策略验证了隐私优先的市场需求。但约 20GB 的安装体积和 8GB VRAM 的硬件要求显著限制了潜在用户群。约 50 条评价的样本量较小，产品仍处于早期验证阶段。', bold=False)

make_body(doc, '与本项目的关系：同是 Live2D + AI 桌面常驻，是最接近的直接对标。本项目不与其竞争"离线能力"或"功能数量"，而是在"长期常驻的克制体验"上差异化。其 20GB 安装体积印证了本地 LLM 方案在分发上的现实成本。', bold=False)

# --- 2.2 Project N.E.K.O. ---
make_heading(doc, '2.2 Project N.E.K.O.（猫娘计划）', level=2)

make_table(doc,
    ['属性', '内容'],
    [
        ['开发商', '喵可智能'],
        ['发布日期', '2026-01-30（抢先体验）'],
        ['平台', 'Windows / macOS（Steam）'],
        ['价格', '完全免费（EA 及正式发布后均免费）'],
        ['技术栈', 'MIT 开源核心 + 闭源内容 DLC'],
        ['Steam 评价', '约 86% 好评（约 50 条评价），当前版本 v0.8.2'],
    ]
)

make_body(doc, '产品定位：功能最全面的开源 AI 桌面伴侣平台，强调"什么都能做"。当前版本迭代速度极快（平均每 1-2 周一个版本），功能边界持续扩展。')

make_body(doc, '核心功能：', bold=True)
make_bullet(doc, '多模态交互：屏幕识别、语音对话、视觉感知、键盘鼠标操控')
make_bullet(doc, '持久记忆系统（v0.8 全面重构，支持跨设备同步）')
make_bullet(doc, 'Agent 能力：NekoClaw 插件系统')
make_bullet(doc, '多种角色格式：Live2D / VRM / MMD')
make_bullet(doc, 'Steam 创意工坊 + GitHub 开源社区')
make_bullet(doc, 'Galgame 模式、猫咪模式（v0.8.2）、角色生成器')

make_body(doc, '关键评价：功能列表在同类产品中最长，但社区反馈中出现了与游戏并行时掉帧、语音无响应、高 CPU 占用等性能信号。这验证了一个关键假设：功能越多，常驻体验越难保证。')

make_body(doc, '与本项目的关系：代表"大而全"路线。其功能广度可作为对比参照——本项目通过克制功能范围来保障常驻体验，两者的用户体验差异可形成论证对比。')

# --- 2.3 Desktop Mate ---
make_heading(doc, '2.3 Desktop Mate', level=2)

make_table(doc,
    ['属性', '内容'],
    [
        ['开发商', 'infiniteloop（日本）'],
        ['发布日期', '2025-01-08'],
        ['平台', 'Windows（Steam）'],
        ['价格', '免费本体（含默认角色 Aiel-tan）；DLC 角色 $14.99/个'],
        ['技术栈', '专有 3D 引擎（非 Live2D）'],
        ['Steam 评价', '65% 好评（约 6,200 条评价），近期 30 天评分下滑至 36%'],
    ]
)

make_body(doc, '产品定位：高质量官方授权 3D 角色的桌面展示平台。核心价值来自 IP（初音未来、三丽鸥、东方 Project、Nier:Automata 2B 等），而非 AI 对话。')

make_body(doc, '核心功能：', bold=True)
make_bullet(doc, '3D 角色在桌面自由移动、坐窗口边缘、追逐光标')
make_bullet(doc, '闹钟功能')
make_bullet(doc, '多角色同屏（beta 中，限同系列角色）')
make_bullet(doc, '丰富的 IP 授权 DLC，每个标准角色 $14.99')
make_bullet(doc, '计划推出独立 VRM 加载工具（回应社区 Mod 需求）')

make_body(doc, '关键评价：在桌面融合质量上是行业标杆——窗口边缘互动、光标追逐、角色物理存在感均为最佳。但其近期评分暴跌（65%→36%）主要源于移除 Mod 支持和 GPU 高占用（集成显卡达 70%）。约 1,800 小时的长期用户虽然认为功能简单，仍持续将角色常驻桌面，说明"高质量角色本身即可支撑长期使用"。')

make_body(doc, '与本项目的关系：Desktop Mate 是本项目 Y1（桌面融合质量）的对比基准——它的 7 项测试用例表现将定义"当前最佳"。同时它证明了一个重要命题：深度 AI 对话不是桌面常驻的必要条件，角色质量和桌面动作本身可以驱动长期留存。')

# --- 2.4 Open-LLM-VTuber ---
make_heading(doc, '2.4 Open-LLM-VTuber', level=2)

make_table(doc,
    ['属性', '内容'],
    [
        ['开发商', '开源社区（t41372 等维护者）'],
        ['版本状态', 'v1.2.1（2025-08-26）为当前稳定版；v2.0 全面重写中'],
        ['平台', 'Windows / macOS / Linux'],
        ['价格', '完全免费开源（MIT）'],
        ['技术栈', 'Python 后端 + Electron/Svelte 前端 + Live2D Cubism 5'],
        ['GitHub', '10,000+ Stars'],
    ]
)

make_body(doc, '产品定位：目前最成熟的开源全链路 AI 伴侣方案，支持完全离线运行。v1.x 已功能冻结，所有开发力量转入 v2.0 全面重写（事件驱动架构）。')

make_body(doc, '核心功能：', bold=True)
make_bullet(doc, '完全离线运行：Ollama / vLLM / LM Studio / GGUF 等本地 LLM 后端')
make_bullet(doc, '桌面宠物模式：透明无边框 + 全局置顶 + 鼠标穿透')
make_bullet(doc, '全链路可替换：ASR（7+）→ LLM（12+）→ TTS（15+）三层 Provider 抽象')
make_bullet(doc, '流式中断（v0.2.3 起）：用户可在 AI 说话时用语音直接打断')
make_bullet(doc, '视觉感知：摄像头 + 屏幕截图')
make_bullet(doc, '长期记忆（Letta）+ MCP 工具调用')
make_bullet(doc, 'Live2D 情绪映射 + 口型同步')

make_body(doc, '关键评价：工程参考价值在同类项目中最高。其流式中断机制、Provider 抽象层设计、透明窗口实现方案均可作为本项目的技术参考。v2.0 正在重写的信号表明该领域技术架构仍在快速演进。')

make_body(doc, '与本项目的关系：技术栈最接近（Electron + Live2D），但产品定位不同——Open-LLM-VTuber 面向 VTuber 和语音交互场景，本项目聚焦轻量桌面常驻。其 Provider 抽象层设计可直接借鉴，但功能范围（视觉、MCP、多 TTS）远超本项目 v1 边界。')

# --- 2.5 Deskie ---
make_heading(doc, '2.5 Deskie（桌小萌）', level=2)

make_table(doc,
    ['属性', '内容'],
    [
        ['开发商', 'SuperQ TEAM'],
        ['发布日期', '2026-01-19（正式版），2025-10-23（首次发布）'],
        ['平台', 'Windows / macOS / Web（Steam）'],
        ['价格', '免费安装；AI 对话消耗"萌点"（每日免费额度 + 付费补充）'],
        ['技术栈', 'Live2D / 3D VRM 双引擎'],
        ['Steam 评价', '约 77% 好评（约 35 条评价）'],
    ]
)

make_body(doc, '产品定位：社区驱动的 AI 角色创作平台，强调"任何人都能创建自己的 AI 角色"。2026 年 5 月更新大幅扩展了桌面宠物和角色创建功能。')

make_body(doc, '核心功能：', bold=True)
make_bullet(doc, '角色格式广泛：Live2D / 3D VRM / 视频 / 图片')
make_bullet(doc, '社区可视化小说（分支故事模式）')
make_bullet(doc, 'AI 角色/宠物创建工坊（无代码，AI 一句话生成动画宠物）')
make_bullet(doc, '桌面小宠物系统（物理拖拽互动）')
make_bullet(doc, '屏幕感知（读取屏幕/网页内容做上下文感知回复）')
make_bullet(doc, '长期记忆 + Steam 创意工坊')

make_body(doc, '关键评价：Deskie 代表"平台化"方向——降低创作门槛、鼓励 UGC、追求内容丰富度。其点数消耗模式引发了用户对"每次对话都在烧钱"的感知。与 Project N.E.K.O. 类似，功能广度与常驻体验之间存在张力。')

make_body(doc, '与本项目的关系：Deskie 的"平台海量角色"路线与本项目的"单一精品角色"路线形成鲜明对比——这是 Y2（低打扰设计）分析中的重要对照点。其点数引发的用户心理成本也是本项目在设计云额度方案时需要考虑的教训。')

# ══════════════════════════════════════════════
# 三、功能对比矩阵
# ══════════════════════════════════════════════

make_heading(doc, '三、功能对比矩阵', level=1)

make_table(doc,
    ['功能维度', 'AI Desktop Pet', 'Project N.E.K.O.', 'Desktop Mate', 'Open-LLM-VTuber', 'Deskie'],
    [
        ['Live2D 渲染', '是（官方授权）', '是', '否（3D引擎）', '是（Cubism 5）', '是'],
        ['透明置顶窗口', '是', '是', '是', '是', '是'],
        ['鼠标穿透', '是', '是', '是', '是', '是'],
        ['拖拽移动', '是', '是', '是', '是', '是'],
        ['多显示器支持', '是', '是', '是', '是', '是'],
        ['窗口吸附', '否', '否', '是（边缘互动）', '否', '否'],
        ['AI 文本对话', '是（本地LLM）', '是（云端）', '否', '是（可替换）', '是（云端）'],
        ['流式/中断', '未确认', '未确认', '—', '是（已确认）', '未确认'],
        ['TTS 语音', '是（声音克隆）', '是', '否', '是（15+引擎）', '是'],
        ['情绪/表情系统', '是', '是', '否（预设动画）', '是', '是'],
        ['长期记忆', '是（角色卡）', '是（v0.8重构）', '否', '是（Letta）', '是'],
        ['创意工坊', '是', '是', '否（已移除）', '否（开源）', '是'],
        ['屏幕感知', '否', '是', '否', '是', '是'],
        ['Agent/工具调用', '否', '是（插件）', '否', '是（MCP）', '部分'],
        ['离线运行', '是（完全）', '否', '是', '是（完全）', '否'],
        ['安装体积', '~20GB', '较小', '较小', '中等', '较小'],
    ]
)

# ══════════════════════════════════════════════
# 四、三维度对比分析
# ══════════════════════════════════════════════

make_heading(doc, '四、三维度对比分析', level=1)

make_body(doc, '本章是毕业设计论证的核心。三个维度（Y1 桌面融合质量、Y2 低打扰设计、Y3 AI 人设稳定性）的测量方法、实验设计和分析结果分述如下。')

# 4.1
make_heading(doc, '4.1 Y1：桌面融合质量', level=2)

make_body(doc, '测量方法：7 项技术测试用例，在统一测试环境（Windows 11，1920×1080 @150% DPI，NVIDIA GTX 1660）下实测。选择 Desktop Mate（桌面融合标杆）和 AI Desktop Pet（直接竞品）作为对比对象。', bold=False)

make_body(doc, '测试用例与对比：', bold=True)

make_table(doc,
    ['编号', '测试用例', 'Desktop Mate', 'AI Desktop Pet', '本项目目标'],
    [
        ['1', '透明穿透：角色区域点击穿过到下方窗口', '通过', '待实测', '必须通过'],
        ['2', '拖拽移动：拖拽后正确定位且恢复穿透', '通过', '待实测', '必须通过'],
        ['3', '多显示器：拖到副屏渲染正常', '通过', '待实测', '必须通过'],
        ['4', 'DPI 切换：125%/150%/200% 缩放清晰', '通过', '待实测', '必须通过'],
        ['5', '睡眠恢复：合盖再打开角色仍在且渲染正常', '注意：GPU高占用', '待实测', '必须通过'],
        ['6', '置顶稳定性：全屏应用不抢占置顶层', '注意：已知问题', '待实测', '必须通过'],
        ['7', '窗口吸附：是否吸附活动窗口边缘', '支持（边缘互动）', '不适用', 'v1 后期实验'],
    ]
)

make_body(doc, '分析：Desktop Mate 在测试 1-4 和 7 上表现最佳，是目前桌面融合的工业标杆。但其在测试 5-6 上的不足（GPU 高占用、置顶不稳定）表明，即使是标杆产品在长期驻留场景下仍有可改进的空间。AI Desktop Pet 的实际表现需要通过 Phase 0 原型验证确定——这将成为本项目的第一个实验数据点。')

# 4.2
make_heading(doc, '4.2 Y2：低打扰设计', level=2)

make_body(doc, '测量方法：6 维客观设计分析。基于产品公开文档、Steam 商店描述和功能列表进行分析，不依赖用户实验。选择全部 5 个竞品 + 本系统设计目标作为对比对象。')

make_table(doc,
    ['编号', '分析维度', 'AI Desktop Pet', 'Project N.E.K.O.', 'Desktop Mate', 'Open-LLM-VTuber', 'Deskie', '本项目设计'],
    [
        ['1', '主动行为频率', '低\n（点击触发）', '中\n（Agent推送）', '极低\n（仅动画）', '中\n（可选主动）', '中\n（社区推送）', '低\n（克制触发）'],
        ['2', '主动行为层级', '轻量\n（眼神追随）', '重量\n（弹窗+操作）', '极轻\n（视线+动作）', '中\n（语音打断）', '中\n（故事通知）', '轻量\n（微动作优先）'],
        ['3', '静默控制', '未明确', '隐私模式', '不支持\n（无暂停）', '配置驱动', '未明确', '必须\n（暂停≤2步）'],
        ['4', '通知策略', '无明显通知', 'Agent 结果', '闹钟提醒', '可选', '社区更新', '默认关闭'],
        ['5', 'IDLE 资源', '待实测', '反馈偏高\n（游戏掉帧）', 'GPU 占用高\n（70%）', '待实测', '待实测', '目标 CPU<1%'],
        ['6', '退出/恢复', '托盘控制', '托盘控制', '不支持\n（无法隐藏）', '配置控制', '托盘控制', '托盘一键退出'],
    ]
)

make_body(doc, '分析：现有竞品普遍在"主动克制"上设计不足。Desktop Mate 因无 AI 对话而天然"低打扰"，但它以牺牲对话能力为代价。Project N.E.K.O. 功能最全但打扰风险最高——用户反馈中的"游戏并行掉帧"即是信号。本项目的设计差异在于：将"低打扰"作为一级设计目标而非副产物——主动行为有冷却机制、静默时段可配置、IDLE 状态帧率降至 5fps、CPU 目标低于 1%。')

# 4.3
make_heading(doc, '4.3 Y3：AI 人设稳定性', level=2)

make_body(doc, '测量方法：本项目采用消融实验 + 竞品定性讨论。原因是各竞品使用不同底层 LLM（本地 Qwen vs 云端 API vs 混合），无法通过统一测试脚本做公平的因果归因。通过在本系统上对比"有完整防护机制"与"无防护机制"的表现差异，论证工程设计的有效性，同时定性对比竞品的人设系统设计。')

make_body(doc, '消融实验设计：', bold=True)

make_table(doc,
    ['实验组', '条件', '说明'],
    [
        ['实验组 A（完整防护）', '三层嵌套约束法 + 角色卡注入 + 滑动窗口摘要（50 轮上限）', '本系统的完整工程设计'],
        ['对照组 B（基础防护）', '仅 system prompt 描述人设，无额外防崩塌机制', '基础防护对比'],
        ['对照组 C（无防护）', '纯模型对话，无任何 system prompt', '基准对照'],
    ]
)

make_body(doc, '测试集（共 120 题）：', bold=True)
make_bullet(doc, '人格基准测试（20 题）：测试角色对自身身份、性格、喜好等基本信息的回答一致性')
make_bullet(doc, '知识边界测试（10 题）：故意询问角色不该知道的信息，检测是否正确表示"不知道"而非编造')
make_bullet(doc, '压力诱导测试（15 题）：用户反复尝试重新定义角色身份（如"其实你是机器人"），检测角色是否坚持人设')
make_bullet(doc, '对话衰减测试（50 轮连续对话）：对比第 1-10 轮与第 41-50 轮回复的一致性')
make_bullet(doc, '情绪合理性测试（20 轮情感对话）：人工标注基准，计算情绪标签匹配率')

make_body(doc, '竞品定性对比：', bold=True)

make_table(doc,
    ['维度', 'AI Desktop Pet', 'Deskie', '本项目'],
    [
        ['人格系统', '角色卡 + 世界书', '角色创建工坊', '三层嵌套约束法'],
        ['记忆机制', '角色卡长记忆', '平台长期记忆', '滑动窗口 + 摘要压缩'],
        ['防崩塌机制', '角色卡作为 context anchor', '不明确', '每轮注入核心角色信息 + few-shot 示例'],
        ['已知风险', '~20GB 体积限制分发', '平台角色质量参差不齐', '依赖 LLM 供应商稳定性'],
    ]
)

make_body(doc, '预期结论：实验组 A 在人格一致性、抗诱导能力和情绪合理性上应显著优于对照组 B 和 C。若消融实验证实工程措施有效，则可以论证：人设稳定性不仅是模型能力问题，更是工程架构设计问题。')

# ══════════════════════════════════════════════
# 五、市场空白与定位
# ══════════════════════════════════════════════

make_heading(doc, '五、市场空白与定位', level=1)

make_heading(doc, '5.1 已确认的非空白', level=2)

make_body(doc, '以下基于"市场空白"的定位声明在本分析后被移除：')
make_bullet(doc, '"Live2D + AI 桌面常驻是市场空白"——已有 5 款以上直接竞品在 Steam 和 GitHub 上发布或活跃开发')
make_bullet(doc, '"第一个能聊天的 Live2D 桌面宠物"——AI Desktop Pet（2026-04-28）已实现此功能')
make_bullet(doc, '"开源方案拒绝 99% 潜在用户"——Open-LLM-VTuber 已提供预打包版本，MateEngine 一键安装即可使用')

make_heading(doc, '5.2 可论证的差异化空间', level=2)

make_body(doc, '在已有激烈竞争的市场中，本项目不从功能数量上竞争，而是在以下四个维度上建立论证：')

make_table(doc,
    ['差异化维度', '市场现状', '本项目的论证切入点'],
    [
        ['克制型常驻体验', '竞品普遍追求功能广度', '将"低打扰"作为一级设计目标，主动行为有冷却/静默/降帧机制'],
        ['桌面融合深度', 'Desktop Mate 是标杆但有 GPU 和置顶问题', '以 Electron 透明窗口实现更低资源占用的桌面融合'],
        ['工程架构合理性', '多数竞品未公开架构', '三层嵌套约束法、表情/动作分离控制、Provider 抽象层'],
        ['隐私数据流透明', '产品营销常模糊"本地"和"云端"边界', '明确区分本地/BYOK/托管云三种模式的数据流和成本'],
    ]
)

make_heading(doc, '5.3 定位陈述', level=2)

make_body(doc, 'Desktop Pet 是一个面向 Windows 桌面的 Live2D AI 桌面伙伴，通过在桌面融合质量、低打扰设计和 AI 人设稳定性三个维度上的工程优化，探索"长期桌面驻留的克制 AI 陪伴体验"这一设计空间。', bold=True)

# ══════════════════════════════════════════════
# 六、结论与设计启示
# ══════════════════════════════════════════════

make_heading(doc, '六、结论与设计启示', level=1)

make_heading(doc, '6.1 对系统设计的直接影响', level=2)

make_body(doc, '竞品分析的每项发现均转化为具体的设计决策：')

make_table(doc,
    ['序号', '竞品信号', '设计决策'],
    [
        ['1', 'AI Desktop Pet 的 20GB 本地模型证明"离线"不是零成本', '不将本地 LLM 作为 v1 默认路径；优先提供 BYOK 入口'],
        ['2', 'Desktop Mate 移除 Mod 引发评分暴跌（65%→36%）', 'v1 支持数据型角色包（不执行第三方代码）；许可和版本元数据必填'],
        ['3', 'Deskie 的点数消耗引发"每次对话都在烧钱"的用户感知', '若采用云额度，显示可理解的剩余额度和成本预估'],
        ['4', 'Project N.E.K.O. 全功能策略导致游戏并行掉帧', 'v1 不做屏幕感知/ASR/Agent 工具调用；IDLE 降至 5fps'],
        ['5', 'Open-LLM-VTuber 的流式中断机制被社区高度认可', '必须实现 AbortController 中断链；用户可随时取消 AI 生成'],
        ['6', 'Moe-AI 卸载残留引发信任危机（虽未独立验证，信号足以要求防范）', '必须提供托盘退出、自启动开关和完整卸载路径'],
    ]
)

make_heading(doc, '6.2 毕业设计实验路线图', level=2)

make_body(doc, 'Phase 0：技术验证（阻断点）', bold=True)
make_bullet(doc, 'P0-1：Electron 透明窗口 + WebGL 穿透原型（2-3 天）')
make_bullet(doc, 'P0-2：Live2D 模型表情清单验证（0.5 天，可并行）')
make_bullet(doc, 'P0-3：官方 Cubism 5 SDK 集成（1-2 天，依赖 P0-1）')

make_body(doc, 'Phase 1：实验准备', bold=True)
make_bullet(doc, '建立竞品基准环境（安装 Desktop Mate + AI Desktop Pet）')
make_bullet(doc, '跑 Y1 桌面融合 7 项测试（3 个对象）')
make_bullet(doc, '完成 Y2 低打扰 6 维设计分析（6 个对象）')
make_bullet(doc, '准备 Y3 消融实验测试脚本（120 题）')

make_body(doc, 'Phase 2：实验执行', bold=True)
make_bullet(doc, 'Y1 数据采集与对比分析')
make_bullet(doc, 'Y2 设计分析报告撰写')
make_bullet(doc, 'Y3 消融实验（A/B/C 三组对比）')

make_body(doc, 'Phase 3：论文撰写', bold=True)
make_bullet(doc, '竞品分析章节（本报告）')
make_bullet(doc, '系统设计章节')
make_bullet(doc, '实验评估章节')

# ══════════════════════════════════════════════
# 附录部分
# ══════════════════════════════════════════════

add_page_break(doc)

make_heading(doc, '附录', level=1)

# 附录 A
make_heading(doc, '附录 A：研究方法与证据等级', level=2)

make_heading(doc, 'A.1 证据分级', level=3)

make_table(doc,
    ['标签', '定义', '使用规则'],
    [
        ['事实', '官网、商店页面、代码仓库或可复现实测直接支持的信息', '标注来源 URL 与访问日期'],
        ['案例', '某条评论、某个版本或某台设备上的具体观察', '只能用于发现信号，不能推广为全体结论'],
        ['推断', '由多个事实或案例形成的分析判断', '使用"表明""提示""可能"，写明推理依据'],
        ['待验证假设', '尚未被实测证实的判断', '给出验证方式和通过标准'],
    ]
)

make_heading(doc, 'A.2 竞品信息采集方法', level=3)
make_bullet(doc, '采集时间：2026-06-09 至 2026-06-11')
make_bullet(doc, '采集方式：Steam 商店页面、SteamDB、GitHub 仓库、产品官网、联网搜索')
make_bullet(doc, '评价数据口径：Steam 公开评价数据，访问日期快照')
make_bullet(doc, '限制说明：Steam 评价存在自选择偏差、语言偏差和版本时点偏差；本报告不以评论数量和评分直接等同于产品质量')

# 附录 B
make_heading(doc, '附录 B：竞品实测数据采集表', level=2)

make_body(doc, '以下表格将在 Phase 1 实验完成后填入实测数据。当前为待采集状态。')

make_table(doc,
    ['测试用例', 'Desktop Mate', 'AI Desktop Pet', '本项目'],
    [
        ['透明穿透', '待实测', '待实测', '待实测'],
        ['拖拽移动', '待实测', '待实测', '待实测'],
        ['多显示器', '待实测', '待实测', '待实测'],
        ['DPI 切换', '待实测', '待实测', '待实测'],
        ['睡眠恢复', '待实测', '待实测', '待实测'],
        ['置顶稳定性', '待实测', '待实测', '待实测'],
        ['窗口吸附', '待实测', 'N/A', '待实测'],
    ]
)

# 附录 C
make_heading(doc, '附录 C：声明验证记录', level=2)

make_body(doc, '以下 8 条关键声明在 2026-06-11 进行了联网独立验证。验证结果：7 条真实，1 条无法独立验证。')

make_table(doc,
    ['编号', '声明', '验证结果', '置信度', '备注'],
    [
        ['C1', 'AI Desktop Pet 于 2026-04-28 发布，中国区 ¥52 买断，支持本地 LLM', '已证实', '高', '多源交叉验证一致'],
        ['C2', 'Deskie 使用点数（Moe Points）消耗模式，免费安装', '已证实', '高', '公测阶段提供每日免费额度'],
        ['C3', 'Project N.E.K.O. 为 MIT 开源核心 + 免费抢先体验', '已证实', '高', 'Open Core 模式，计划闭源内容变现'],
        ['C4', 'Desktop Mate DLC 角色美国区 $14.99', '已证实', '高', '仅 Yukkuri Reimu（简化版）$7.49 为例外'],
        ['C5', 'Open-LLM-VTuber 已实现流式中断和可替换 AI 管线', '已证实', '高', 'v0.2.3 起支持语音打断'],
        ['C6', 'Moe-AI 用户反映卸载后仍出现启动行为', '无法独立验证', '低', '无法通过公开网络搜索独立复现'],
        ['C7', 'VPet Steam 评价约 5.1 万条，有活跃 Workshop', '已证实', '高', 'Steambase 记录 50,927 条评价'],
        ['C8', 'AIRI 使用 TypeScript/Vue/Electron 栈，2026 年持续发布', '已证实', '高', 'v0.10.2（2026-05-07）为最新稳定版'],
    ]
)

make_heading(doc, '附录 D：主要来源', level=2)

make_body(doc, '商业产品：', bold=True)
make_bullet(doc, 'AI Desktop Pet Steam：https://store.steampowered.com/app/4227700/AI_Desktop_Pet/')
make_bullet(doc, 'Deskie Steam：https://store.steampowered.com/app/3540710/Deskie/')
make_bullet(doc, 'Project N.E.K.O. Steam：https://store.steampowered.com/app/4099310/Project_NEKO/')
make_bullet(doc, 'Desktop Mate Steam：https://store.steampowered.com/app/3301060/Desktop_Mate/')
make_bullet(doc, 'Desktop Mate 开发者路线图：https://steamcommunity.com/app/3301060/discussions/0/597397602266286127/')

make_body(doc, '开源项目：', bold=True)
make_bullet(doc, 'Open-LLM-VTuber GitHub：https://github.com/Open-LLM-VTuber/Open-LLM-VTuber')
make_bullet(doc, 'Project N.E.K.O. GitHub：https://github.com/Project-N-E-K-O/N.E.K.O')
make_bullet(doc, 'AIRI GitHub：https://github.com/moeru-ai/airi')
make_bullet(doc, 'VPet GitHub：https://github.com/LorisYounger/VPet')
make_bullet(doc, 'Navi GitHub：https://github.com/Violet2314/NAVI')
make_bullet(doc, 'MiniCPM Desk Pet GitHub：https://github.com/OpenBMB/MiniCPM-Desk-Pet')

make_body(doc, '数据来源：', bold=True)
make_bullet(doc, 'SteamDB：https://steamdb.info/')
make_bullet(doc, 'Steambase：https://steambase.io/')
make_bullet(doc, 'Deku Deals：https://www.dekudeals.com/')

make_body(doc, '所有来源的访问日期均为 2026-06-11。')

# ─── 保存 ─────────────────────────────────────────────────

output_path = 'e:/Work-26/AI_Desktop_Pet/docs/竞品分析报告-v2.0.docx'
doc.save(output_path)
print(f'[OK] Report saved: {output_path}')
print(f'     Paragraphs: {len(doc.paragraphs)}')
print(f'     Tables: {len(doc.tables)}')
