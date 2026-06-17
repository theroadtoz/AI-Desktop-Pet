#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
为竞品分析报告添加 Word 自动目录
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def insert_toc_field(paragraph, levels="1-3"):
    """在段落中插入 Word TOC 域代码"""
    # 清除段落现有内容
    for run in paragraph.runs:
        run._r.getparent().remove(run._r)

    run = paragraph.add_run()

    # TOC 域开始标记
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    # TOC 指令
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' TOC \\o "{levels}" \\h \\z \\u '
    run._r.append(instrText)

    # 域分隔符
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar_separate)

    return fldChar_separate, run

def insert_toc_end(paragraph):
    """在段落末尾插入 TOC 域结束标记"""
    run = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

def add_toc_instruction(paragraph):
    """添加提示文字"""
    run = paragraph.add_run('（打开文档后，右键此处 → 更新域 → 更新整个目录）')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def main():
    input_path = r'e:\Work-26\FinalProject\docs\reports\竞品分析报告.docx'

    doc = Document(input_path)

    # 找到目录页（第二个分页符之后的第一个段落是"目  录"标题）
    # 遍历段落找"目  录"
    found = False
    for i, para in enumerate(doc.paragraphs):
        if '目' in para.text and '录' in para.text and para.runs:
            # 确认这是目录页的标题
            if any('目' in r.text for r in para.runs):
                found = True
                # 下一段是提示文字，替换它
                for j in range(i+1, min(i+5, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    # 清空提示段落
                    for r in next_para.runs:
                        r._r.getparent().remove(r._r)

                    # 插入 TOC 域
                    insert_toc_field(next_para, levels="1-3")
                    add_toc_instruction(next_para)
                    insert_toc_end(next_para)
                    break
                break

    if not found:
        print('Warning: TOC heading not found, inserting at start after cover')
        # Fallback: insert TOC after first page break
        for i, para in enumerate(doc.paragraphs):
            if i > 3:  # Skip cover page
                # Insert TOC title
                toc_title = doc.paragraphs[i]
                # Clear
                for r in toc_title.runs:
                    r._r.getparent().remove(r._r)
                # Don't modify this paragraph, add new ones
                break

    doc.save(input_path)
    print(f'TOC field inserted into: {input_path}')
    print('Open in Word, right-click TOC area, select "Update Field" to generate.')

if __name__ == '__main__':
    main()
